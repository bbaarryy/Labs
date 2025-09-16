import matplotlib.pyplot as plt
#import seaborn as sns
from docx import Document

file = open("1.1.4/data.txt", "r")

def get_arr(raw, q):
    i = 0
    arr = []
    summ = 0
    while(i < len(raw)):
        for j in range(q):
            summ += raw[i]
            i+=1
        arr.append(summ)
        summ = 0
    return arr

raw = []
for line in file:
    raw.append(int(line))
print(len(raw))

def print_table(arr, ln, doc):
    for e in range(0,len(arr),ln):
        table = doc.add_table(rows=3, cols=ln+1)
        table.cell(0, 0).text = "Число импульсов"
        table.cell(1, 0).text = "Число случаев"
        table.cell(2, 0).text = "Доля случаев"
        for i in range(e,min(e + ln, len(arr))):
            table.cell(0,i-e+1).text = str(i)
            table.cell(1, i-e+1).text = str(arr[i])
            table.cell(2, i-e+1).text = str(arr[i]/all)
        table.style = 'Light Shading Accent 1'
    
# ans = 0
# for i in range(len(arr20)):
#     if(24.36 - 2*5.198 <= arr20[i] <= 24.36 + 2*5.198):
#         ans += 1
# print(ans)




doc = Document()

X = 20
arr20 = get_arr(raw,X)
table = doc.add_table(rows=X+1, cols=11)

table.cell(0,0).text = "Опыт"
for i in range(0,X):
    table.cell(i+1,0).text = str(i*10)

for i in range(0,10):
    table.cell(0,i+1).text = str(i+1)


for x in range(0,X):
    for y in range(0,10):
        table.cell(x+1,y+1).text = str(arr20[x*10+y])

table.style = 'Light Shading Accent 1'
doc.save('demo.docx')



# arr20 = get_arr(raw,80)
# mx = max(arr20)

# all = 0

# x = [1]
# for i in range(2,mx+1):
#     x.append(i)
# y = [0] * mx

# for i in range(len(arr20)):
#     if(arr20[i] != 0):
#         y[arr20[i]-1]+=1
#         all+=1

# print_table(y,13,doc)


#doc.save('demo.docx')

# for i in range(len(y)):
#     y[i] = y[i]/all

# ax = plt.gca()
# #ax.axes.xaxis.set_visible(False)
# vis = [""]
# # e = [""]
# for i in range(0,mx,10):
#     for j in range(9):
#         vis.append("")
#     vis.append(str(i+10))
# ax.set_xticklabels(vis)
# sns.barplot(x=x, y=y, color='blue')
# plt.xlabel("n") # Подпись оси X
# plt.ylabel("w") # Подпись оси Y
# plt.title('Гистограмма для t = 80') #Название


# plt.show()